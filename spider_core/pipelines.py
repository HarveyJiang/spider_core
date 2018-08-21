# -*- coding: utf-8 -*-

# Define your item pipelines here
#
# Don't forget to add your pipeline to the ITEM_PIPELINES setting
# See: http://doc.scrapy.org/en/latest/topics/item-pipeline.html
import re

import scrapy
from scrapy.pipelines.files import FilesPipeline
from scrapy.pipelines.images import ImagesPipeline


class SpiderCorePipeline(object):
    def process_item(self, item, spider):
        return item


class SpiderCoreMongoDBPipeline(object):

    def process_item(self, item, spider):
        spider.mongod_helper.insert_documents(item, spider.spider_name)
        return item


class SpiderCoreSaveWordPipeline(object):
    def process_item(self, item, spider):
        from docx import Document
        import html2text
        document = Document()
        document.add_heading(item['title'], 0)
        mk = html2text.HTML2Text()
        mk.ignore_links = True
        content = mk.handle(item['content'])
        document.add_paragraph(content)
        title = re.sub('\W+', '', item['title'])
        document.save('./docs/{}.docx'.format(title))
        print('save... word')
        return item


class SpiderCoreImagesPipeline(ImagesPipeline):
    def get_media_requests(self, item, info):
        user_agent = 'Mozilla/5.0 (Windows NT 6.3; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/45.0.2454.101 Safari/537.36'
        for image_url in item.get('image_urls'):
            # print('image_url', image_url)
            # fn= response.url.split('/')[-1]
            # fsock = open(r"F:\UserData\Desktop\\"+fn, "wb")
            # fsock.write(response.body)
            # print 'image_url:',image_url
            yield scrapy.Request(image_url, meta=item, headers={'User-Agent': user_agent})

    def file_path(self, request, response=None, info=None):
        image_name = request.url.split('/')[-1]
        # print('url2', response.url)
        # image_name = re.sub('\W+', '', request.meta.get('title')) + str(uuid.uuid1())[:8]
        # item = request.meta['item']  # Like this you can use all from item, not just url.
        # image_name = item.get('title')
        return 'full/{0}/{1}'.format(request.meta.get('title'), image_name)

    def item_completed(self, results, item, info):
        # image_paths = [x['path'] for ok, x in results if ok]
        # if not image_paths:
        #     raise DropItem("Item contains no images")
        # item['image_paths'] = image_paths
        # item['image_paths'] = results
        print('image over')
        return item


class SpiderCoreFilesPipeline(FilesPipeline):
    def get_media_requests(self, item, info):
        for file_url in item['file_urls']:
            yield scrapy.Request(file_url)
